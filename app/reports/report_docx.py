from docx import Document


def generate_report_docx(report: dict, file_path: str):
    doc = Document()

    # Title
    doc.add_heading("Assessment Report", level=1)

    # Summary
    doc.add_heading("Summary", level=2)
    for line in report["summary"]:
        doc.add_paragraph(line)

    # Scores
    doc.add_heading("Overall Score", level=2)
    scores = report["scores"]
    doc.add_paragraph(
        f"Score: {scores['total_score']} / {scores['max_score']} "
        f"({scores['percentage']}%)"
    )
    doc.add_paragraph(f"Level: {scores['level']}")
    doc.add_paragraph(f"Status: {scores['status']}")

    # Pros
    doc.add_heading("Strengths (Pros)", level=2)
    for p in report["pros"]:
        doc.add_paragraph(p, style="List Bullet")

    # Cons
    doc.add_heading("Weaknesses (Cons)", level=2)
    for c in report["cons"]:
        doc.add_paragraph(c, style="List Bullet")

    # Scope of Improvement
    doc.add_heading("Scope of Improvement", level=2)
    for s in report["scope_of_improvement"]:
        doc.add_paragraph(s, style="List Number")

    doc.save(file_path)
